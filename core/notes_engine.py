import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration
FONT_NAME       = "Nirmala UI"
FONT_TITLE      = 18
FONT_HEADING    = 16
FONT_SUBHEADING = 15
FONT_BODY       = 14
FONT_SUBBULLET  = 13
FONT_NOTE       = 13
LINE_SPACING    = 1.5

COL_TITLE      = "1F3864"
COL_HEADING    = "2E4057"
COL_SUBHEADING = "1F4E79"
COL_SUB        = "595959"
COL_NOTE       = "7F7F7F"
COL_ACCENT     = "4472C4"
COL_RULE       = "BDD7EE"

def hex_color(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def apply_spacing(para, line_spacing=LINE_SPACING, before=0, after=6):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = line_spacing
    fmt.space_before      = Pt(before)
    fmt.space_after       = Pt(after)

def apply_font(run, size=FONT_BODY, bold=False, italic=False, color=None):
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = hex_color(color)
    rPr    = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"),    FONT_NAME)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def add_run(para, text, size=FONT_BODY, bold=False, italic=False, color=None):
    run = para.add_run(text)
    apply_font(run, size=size, bold=bold, italic=italic, color=color)
    return run

def add_border(para, side="bottom", style="single", size=6, color=COL_ACCENT, space=4):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el   = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   style)
    el.set(qn("w:sz"),    str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)

def safe_add_paragraph(doc, text="", style=None, bullet=False):
    """Adds a paragraph with a style, falling back to manual formatting if style is missing."""
    try:
        p = doc.add_paragraph(style=style)
    except (ValueError, KeyError):
        p = doc.add_paragraph()
        if bullet:
            p.add_run("• ")
    
    if text:
        p.add_run(text)
    return p

def extract_english_title(title):
    match = re.search(r'\((.*?)\)', title)
    if match:
        return match.group(1).strip()
    return title.strip()

def build_notes_document(notes_data, output_path, template_path=None):
    doc = Document(template_path) if template_path else Document()

    # Layout Setup
    for sec in doc.sections:
        sec.page_width         = Inches(8.27)
        sec.page_height        = Inches(11.69)
        sec.top_margin         = Inches(1.0)
        sec.bottom_margin      = Inches(1.0)
        sec.left_margin        = Inches(1.0)
        sec.right_margin       = Inches(1.0)

    # Main title
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(t_para, notes_data.get("title", "Notes"), size=FONT_TITLE, bold=True, color=COL_TITLE)
    add_border(t_para, side="bottom", style="single", size=8, color=COL_ACCENT, space=6)
    apply_spacing(t_para, line_spacing=LINE_SPACING, before=0, after=14)

    # Sections
    for idx, section in enumerate(notes_data.get("sections", [])):
        if idx > 0:
            gap = doc.add_paragraph()
            apply_spacing(gap, line_spacing=1.0, before=0, after=10)

        h_para = doc.add_paragraph()
        add_run(h_para, section.get("heading", ""), size=FONT_HEADING, bold=True, color=COL_HEADING)
        add_border(h_para, side="bottom", style="single", size=4, color=COL_RULE, space=2)
        apply_spacing(h_para, line_spacing=LINE_SPACING, before=6, after=4)

        for item in section.get("bullets", []):
            if isinstance(item, str):
                item = {"type": "bullet", "text": item, "sub_bullets": []}

            item_type = item.get("type", "bullet")
            text      = item.get("text", "")

            if item_type == "sub_heading":
                sh_para = doc.add_paragraph()
                add_run(sh_para, text, size=FONT_SUBHEADING, bold=True, color=COL_SUBHEADING)
                apply_spacing(sh_para, line_spacing=LINE_SPACING, before=6, after=2)

            elif item_type == "note":
                n_para = doc.add_paragraph()
                n_para.paragraph_format.left_indent = Inches(0.3)
                add_run(n_para, text, size=FONT_NOTE, italic=True, color=COL_NOTE)
                apply_spacing(n_para, line_spacing=LINE_SPACING, before=4, after=4)

            elif item_type == "numbered":
                n_para = safe_add_paragraph(doc, style="List Number")
                add_run(n_para, text, size=FONT_BODY)
                apply_spacing(n_para, line_spacing=LINE_SPACING, before=0, after=4)
                for sub in item.get("sub_bullets", []):
                    sb_para = safe_add_paragraph(doc, style="List Bullet 2", bullet=True)
                    add_run(sb_para, sub, size=FONT_SUBBULLET, color=COL_SUB)
                    apply_spacing(sb_para, line_spacing=LINE_SPACING, before=0, after=2)

            elif item_type == "table":
                headers = item.get("headers", [])
                rows    = item.get("rows", [])
                if not headers: continue
                pre = doc.add_paragraph()
                apply_spacing(pre, line_spacing=1.0, before=0, after=4)
                tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = "Table Grid"
                col_width = Inches(6.27 / len(headers))
                
                # Header
                for i, hdr_text in enumerate(headers):
                    cell = tbl.rows[0].cells[i]
                    cell.width = col_width
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), COL_HEADING)
                    tcPr.append(shd)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].add_run(hdr_text)
                    apply_font(run, size=FONT_BODY - 1, bold=True, color="FFFFFF")
                
                # Data Rows
                for r_idx, row_data in enumerate(rows):
                    fill = "DEEAF1" if r_idx % 2 == 0 else "FFFFFF"
                    for c_idx, cell_text in enumerate(row_data):
                        cell = tbl.rows[r_idx + 1].cells[c_idx]
                        cell.width = col_width
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
                        tcPr.append(shd)
                        run = cell.paragraphs[0].add_run(cell_text)
                        apply_font(run, size=FONT_BODY - 1)

            else:
                b_para = safe_add_paragraph(doc, style="List Bullet", bullet=True)
                add_run(b_para, text, size=FONT_BODY)
                apply_spacing(b_para, line_spacing=LINE_SPACING, before=0, after=4)
                for sub in item.get("sub_bullets", []):
                    sb_para = safe_add_paragraph(doc, style="List Bullet 2", bullet=True)
                    add_run(sb_para, sub, size=FONT_SUBBULLET, color=COL_SUB)
                    apply_spacing(sb_para, line_spacing=LINE_SPACING, before=0, after=2)

    doc.save(output_path)
