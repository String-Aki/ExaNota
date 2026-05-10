import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COL_NAVY = (0x1F, 0x49, 0x7D)

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_question(doc, number, q_text, options, correct_letter, is_tamil, is_teacher_copy):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(10.5)
    run_num.font.color.rgb = RGBColor(*COL_NAVY)

    run_q = p.add_run(q_text)
    run_q.font.size = Pt(10.5)

    letters = ['அ)', 'ஆ)', 'இ)', 'ஈ)'] if is_tamil else ['A)', 'B)', 'C)', 'D)']
    opt_keys = ['A', 'B', 'C', 'D']
    
    for i, (letter_key, opt_text) in enumerate(zip(opt_keys, options)):
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Inches(0.35)
        op.paragraph_format.space_before = Pt(1)
        op.paragraph_format.space_after = Pt(1)
        
        is_correct = (letter_key == correct_letter)
        should_highlight = is_correct and is_teacher_copy
        
        run_l = op.add_run(f"  {letters[i]}  ")
        run_l.font.size = Pt(10)
        if should_highlight:
            run_l.bold = True
            run_l.font.color.rgb = RGBColor(*COL_NAVY)
            
        run_o = op.add_run(str(opt_text))
        run_o.font.size = Pt(10)
        if should_highlight:
            run_o.bold = True
            run_o.font.color.rgb = RGBColor(*COL_NAVY)

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    return doc

def generate_exam_docx(questions_data, is_tamil=True, is_teacher_copy=False):
    doc = setup_document()
    
    # Check if input is the full dictionary or just a list
    if isinstance(questions_data, dict):
        title = questions_data.get("title", "வினாத்தாள் (Question Paper)")
        sections = questions_data.get("sections", [])
        # Flatten sections into a list of questions with section markers
        flat_questions = []
        for s in sections:
            heading = s.get("heading")
            qs = s.get("questions", [])
            if qs and heading:
                qs[0]["section_title"] = heading
            flat_questions.extend(qs)
    else:
        title = "வினாத்தாள் (Question Paper)"
        flat_questions = questions_data

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*COL_NAVY)
    
    # Instructions
    p_inst = doc.add_paragraph()
    r_inst = p_inst.add_run("வழிமுறைகள் (Instructions):")
    r_inst.bold = True; r_inst.font.size = Pt(11); r_inst.font.color.rgb = RGBColor(*COL_NAVY)

    instructions = [
        "1. ஒவ்வொரு வினாவிற்கும் சரியான விடையை மட்டும் தேர்வு செய்யவும்.",
        "2. ஒவ்வொரு சரியான விடைக்கும் 1 மதிப்பெண் வழங்கப்படும்.",
        f"3. மொத்த வினாக்கள்: {len(flat_questions)}  |  மொத்த மதிப்பெண்கள்: {len(flat_questions)}"
    ]
    for inst in instructions:
        pi = doc.add_paragraph()
        pi.paragraph_format.left_indent = Inches(0.25)
        run = pi.add_run(inst); run.font.size = Pt(10)

    add_rule(doc)
    doc.add_paragraph()

    for idx, item in enumerate(flat_questions):
        if "section_title" in item:
            doc.add_paragraph()
            sh = doc.add_paragraph()
            sh_run = sh.add_run(f" {item['section_title']}")
            sh_run.bold = True; sh_run.font.size = Pt(11); sh_run.font.color.rgb = RGBColor(*COL_NAVY)
            pPr = sh._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:color'), '2E74B5')
            pBdr.append(bottom); pPr.append(pBdr)

        add_question(doc, idx + 1, item.get("q", ""), item.get("opts", []), item.get("ans", ""), is_tamil, is_teacher_copy)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_answer_key_docx(questions_data):
    doc = setup_document()
    
    if isinstance(questions_data, dict):
        flat_questions = []
        for s in questions_data.get("sections", []):
            flat_questions.extend(s.get("questions", []))
    else:
        flat_questions = questions_data

    ah = doc.add_paragraph()
    ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ah_run = ah.add_run("விடைகள் (Answer Key)")
    ah_run.bold = True; ah_run.font.size = Pt(14); ah_run.font.color.rgb = RGBColor(*COL_NAVY)
    add_rule(doc)
    doc.add_paragraph()

    total_q = len(flat_questions)
    cols = 10
    rows = (total_q + cols - 1) // cols
    tbl = doc.add_table(rows=rows + 1, cols=cols)
    tbl.style = 'Table Grid'

    for c in range(cols):
        tbl.rows[0].cells[c].paragraphs[0].add_run("வி.எண்").bold = True

    for idx, q in enumerate(flat_questions):
        row, col = idx // cols + 1, idx % cols
        cell = tbl.rows[row].cells[col]
        run = cell.paragraphs[0].add_run(f"{idx+1}. {q.get('ans', '')}")
        run.bold = True; run.font.color.rgb = RGBColor(*COL_NAVY)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
